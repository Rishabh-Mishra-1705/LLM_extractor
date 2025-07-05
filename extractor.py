from openai import OpenAI
from docx import Document
import json
import re


def read_docx(file_path):
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            full_text.append(" | ".join(row_text))

    return "\n\n".join(full_text)


def create_prompt(doc_text, fields):
    field_descriptions = {
        "Business Source": "Source type (Direct/Broker/Not Found) from header",
        "Original Insured Name": "Name of original insured party",
        "Broker Name": "Name of broker if applicable",
        "Reinsured Name": "Name of reinsured company",
        "Insured Address": "Physical address of insured",
        "Type of Contract": "Contract type description",
        "Policy Duration": "Duration in years or months",
        "Validity From": "Start date (DD/MM/YYYY format)",
        "Validity To": "End date (DD/MM/YYYY format)",
        "Currency": "Currency code (e.g. MXN)",
        "Exchange Rate": "Currency exchange rate if available",
        "Territorial Limit": "Geographical coverage description",
        "Reinsurance Scheme": "Type of reinsurance scheme",
        "Maximum Contract Limit": "Maximum coverage amount with currency",
        "Number of Insured": "Total number of insured individuals",
        "Total Sum Insured": "Total insured amount with currency",
        "Reinsurance Commission": "Commission percentage for reinsurer",
        "Premium Payment Frequency": "How often premiums are paid",
        "Coverage Types": "List of coverage types included",
        "Special Clauses": "List of special clauses mentioned"
    }

    instructions = (
        "Extract EXACT values from the insurance document for these fields:\n"
        "1. Return ONLY valid JSON with the specified keys\n"
        "2. Use NULL for missing fields\n"
        "3. Preserve original language and formatting\n"
        "4. Handle multi-line values appropriately\n"
        "5. Extract dates in DD/MM/YYYY format\n"
        "6. For monetary values, include currency code\n"
        "7. For lists, return as JSON arrays\n\n"
        "FIELD DESCRIPTIONS:\n"
    )

    fields_str = "\n".join(
        f"- {field}: {field_descriptions.get(field, '')}"
        for field in fields
    )

    examples = (
        "\n\nEXAMPLE OUTPUT:\n"
        '{"Original Insured Name": "Instituto de Servicios Educativos y Pedagógicos de Baja California", '
        '"Validity From": "01/02/2024", '
        '"Validity To": "31/12/2024", '
        '"Currency": "MXN", '
        '"Number of Insured": "31,175", '
        '"Reinsurance Commission": "1.5%", '
        '"Premium Payment Frequency": "Trimestral", '
        '"Coverage Types": ["Fallecimiento", "Invalidez total y permanente", "Anticipo de Suma Asegurada por enfermedades terminales", "Anticipo de Suma asegurada por Gastos Funerarios", "Exención de pago de prima por licencia médica"]}\n\n'
        "DOCUMENT TEXT:\n"
    )

    return instructions + fields_str + examples + doc_text

fields_to_extract = [
    "Business Source", "Original Insured Name", "Broker Name",
    "Reinsured Name", "Insured Address", "Type of Contract",
    "Policy Duration", "Validity From", "Validity To", "Currency",
    "Exchange Rate", "Territorial Limit", "Reinsurance Scheme",
    "Maximum Contract Limit", "Number of Insured", "Total Sum Insured",
    "Reinsurance Commission", "Premium Payment Frequency",
    "Coverage Types", "Special Clauses"
]

# Initialize client
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key="your-api-token",
)


def extract_fields_from_doc(docx_path):
    doc_text = read_docx(docx_path)
    prompt = create_prompt(doc_text, fields_to_extract)

    response = client.chat.completions.create(
        model="deepseek/deepseek-r1-0528:free",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=2000,
        response_format={"type": "json_object"}
    )

    result = response.choices[0].message.content
    print("\n--- Extracted Fields ---\n")

    try:
        parsed = json.loads(result)
        print(json.dumps(parsed, indent=2, ensure_ascii=False))
        return parsed
    except json.JSONDecodeError:
        print("Invalid JSON format received:")
        print(result)
        return {"error": "Failed to parse JSON output"}


if __name__ == "__main__":
    result = extract_fields_from_doc("SkyRe - SEP de B.C. - Slip de Reaseguro (O).docx")
